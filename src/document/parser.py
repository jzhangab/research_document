from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional


@dataclass
class ParsedDocument:
    raw_text: str
    page_map: dict[int, str]
    file_name: str
    file_type: str
    metadata: dict = field(default_factory=dict)


class DocumentParser:
    def parse(self, file_path: str | Path) -> ParsedDocument:
        path = Path(file_path)
        ext = path.suffix.lower().lstrip(".")
        if ext == "pdf":
            return self._parse_pdf(path)
        elif ext in ("docx", "doc"):
            return self._parse_docx(path)
        else:
            return self._parse_txt(path)

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        page_map: dict[int, str] = {}
        for i, page in enumerate(doc):
            page_map[i] = page.get_text()

        meta = doc.metadata or {}
        doc.close()

        return ParsedDocument(
            raw_text="\n".join(page_map.values()),
            page_map=page_map,
            file_name=path.name,
            file_type="pdf",
            metadata={
                "title": meta.get("title", ""),
                "author": meta.get("author", ""),
                "creation_date": meta.get("creationDate", ""),
            },
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        sections: dict[int, str] = {}
        current_section = 0
        buffer: list[str] = []

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                if buffer:
                    sections[current_section] = "\n".join(buffer)
                    current_section += 1
                    buffer = []
            buffer.append(para.text)

        if buffer:
            sections[current_section] = "\n".join(buffer)

        props = doc.core_properties
        return ParsedDocument(
            raw_text="\n".join(sections.values()),
            page_map=sections,
            file_name=path.name,
            file_type="docx",
            metadata={
                "title": props.title or "",
                "author": props.author or "",
                "creation_date": str(props.created or ""),
            },
        )

    def _parse_txt(self, path: Path) -> ParsedDocument:
        text = path.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(
            raw_text=text,
            page_map={0: text},
            file_name=path.name,
            file_type="txt",
        )
